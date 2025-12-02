"""
DOCX Extractor using python-docx
Extracts text content and metadata from Word documents
"""

import logging
from pathlib import Path
from typing import List

from .base import BaseExtractor, ExtractedDocument, ExtractionError

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """
    Extracts text from DOCX files using python-docx.

    Features:
    - Full text extraction with paragraph structure
    - Preserves heading hierarchy (converted to markdown)
    - Extracts document properties (title, author, created date)
    - Handles tables (converted to text)

    Limitations:
    - Images are skipped (text only)
    - Complex formatting lost
    - .doc (old format) not supported - only .docx
    """

    SUPPORTED_EXTENSIONS = ['.docx']

    def extract(self, file_path: Path) -> ExtractedDocument:
        """Extract text and metadata from DOCX"""
        self._validate_file(file_path)

        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            raise ExtractionError("python-docx not installed. Run: pip install python-docx")

        warnings = []
        text_parts = []

        try:
            doc = Document(str(file_path))
        except PackageNotFoundError:
            raise ExtractionError(f"Invalid or corrupted DOCX file: {file_path.name}")
        except Exception as e:
            if "password" in str(e).lower() or "encrypted" in str(e).lower():
                raise ExtractionError(f"Document is password-protected: {file_path.name}")
            raise ExtractionError(f"Failed to open DOCX: {e}")

        # Extract metadata from core properties
        title = None
        author = None
        creation_date = None

        try:
            core_props = doc.core_properties
            title = core_props.title if core_props.title else None
            author = core_props.author if core_props.author else None
            if core_props.created:
                creation_date = core_props.created.isoformat()
        except Exception as e:
            logger.debug(f"Could not extract DOCX metadata: {e}")

        # Extract paragraphs with heading detection
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Convert headings to markdown format for better structure
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.replace('Heading ', '').replace('Heading', '1'))
                    text = '#' * level + ' ' + text
                except ValueError:
                    pass

            text_parts.append(text)

        # Extract tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(f"\n[Table {table_count}]\n{table_text}")

        if table_count > 0:
            logger.debug(f"Extracted {table_count} tables from {file_path.name}")

        # Check for images (we can't extract them, but note their presence)
        try:
            image_count = len(doc.inline_shapes)
            if image_count > 0:
                warnings.append(f"{image_count} image(s) skipped (text extraction only)")
        except Exception:
            pass

        content = "\n\n".join(text_parts).strip()

        logger.info(f"Extracted {len(content)} chars from {file_path.name}")

        return ExtractedDocument(
            content=content,
            format_type='docx',
            title=title if title and title.strip() else None,
            author=author if author and author.strip() else None,
            creation_date=creation_date,
            extraction_warnings=warnings,
            original_path=str(file_path),
            extra_metadata={
                'paragraph_count': len(doc.paragraphs),
                'table_count': table_count,
            }
        )

    def _extract_table(self, table) -> str:
        """Convert a DOCX table to text format"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        if not rows:
            return ""

        # Add separator after header row
        if len(rows) > 1:
            header = rows[0]
            separator = " | ".join(["---"] * len(rows[0].split(" | ")))
            return "\n".join([header, separator] + rows[1:])

        return "\n".join(rows)
